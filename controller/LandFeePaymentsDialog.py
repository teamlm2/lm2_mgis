__author__ = 'mwagner'
# coding=utf8

from ..view.Ui_LandFeePaymentsDialog import *
from ..utils.PluginUtils import *
from ..model.BsPerson import *
from ..model.CtFee import *
from ..model.DatabaseHelper import *
from sqlalchemy.sql import func
from ..model.SetPayment import *
from datetime import datetime
from docx import Document, shape
from docx.shared import Inches, Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from ..model.FeeUnified import *
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PAYMENT_DATE = 0
AMOUNT_PAID = 1
PAYMENT_TYPE = 2


class LandFeePaymentsDialog(QDialog, Ui_LandFeePaymentsDialog, DatabaseHelper):

    def __init__(self, contract, parent=None):

        super(LandFeePaymentsDialog,  self).__init__(parent)
        DatabaseHelper.__init__(self)
        self.setupUi(self)
        self.__setup_twidgets()
        self.close_button.clicked.connect(self.reject)
        self.contract = contract
        self.session = SessionHandler().session_instance()
        self.payment_date_edit.setDate(QDate.currentDate())
        self.amount_paid_sbox.setMinimum(0)
        self.amount_paid_sbox.setMaximum(25000000)
        self.amount_paid_sbox.setSingleStep(1000)
        self.fine_payment_date_edit.setDate(QDate.currentDate())
        self.fine_amount_paid_sbox.setMinimum(0)
        self.fine_amount_paid_sbox.setMaximum(25000000)
        self.fine_amount_paid_sbox.setSingleStep(200)
        self.year_sbox.setMinimum(1950)
        self.year_sbox.setMaximum(2200)
        self.year_sbox.setSingleStep(1)
        self.year_sbox.setValue(QDate.currentDate().year())
        self.status_label.clear()

        self.__load_data()

    def __setup_twidgets(self):

        self.payment_twidget.setAlternatingRowColors(True)
        self.payment_twidget.setSelectionBehavior(QTableWidget.SelectRows)
        self.payment_twidget.setSelectionMode(QTableWidget.SingleSelection)

        self.fine_payment_twidget.setAlternatingRowColors(True)
        self.fine_payment_twidget.setSelectionBehavior(QTableWidget.SelectRows)
        self.fine_payment_twidget.setSelectionMode(QTableWidget.SingleSelection)

    def __load_data(self):

        self.begin_date.setDate(self.contract.contract_begin)
        self.end_date.setDate(self.contract.contract_end)
        self.contract_number_edit.setText(self.contract.contract_no)
        begin = PluginUtils.convert_python_date_to_qt(self.contract.contract_begin)
        self.contract_begin_edit.setText(begin.toString(Constants.DATE_FORMAT))
        end = PluginUtils.convert_python_date_to_qt(self.contract.contract_end)
        self.contract_end_edit.setText(end.toString(Constants.DATE_FORMAT))
        self.contract_status_edit.setText(self.contract.status_ref.description)

        self.__populate_payment_type_cboxes()
        self.__populate_contractor_cbox()

    def __populate_payment_type_cboxes(self):

        self.__populate_payment_type_cbox(self.payment_type_cbox)
        self.__populate_payment_type_cbox(self.fine_payment_type_cbox)

    def __populate_payment_type_cbox(self, cbox):

        for payment_type in self.session.query(ClPaymentType).order_by(ClPaymentType.code).all():

            cbox.addItem(u"{0}".format(payment_type.description), payment_type.code)

    def __populate_contractor_cbox(self):

        for fee in self.contract.fees:
            print fee.person
            person_id = fee.person
            for name, first_name in self.session.query(BsPerson.name, BsPerson.first_name)\
                    .filter(BsPerson.person_id == person_id):
                if first_name is None:
                    self.select_contractor_cbox.addItem(u"{0}".format(name), person_id)
                else:
                    self.select_contractor_cbox.addItem(u"{0}, {1}".format(name, first_name), person_id)

    @pyqtSlot(int)
    def on_select_contractor_cbox_currentIndexChanged(self, idx):

        self.__clear_controls()

        if idx == -1:
            return

        person_id = self.select_contractor_cbox.itemData(idx, Qt.UserRole)
        self.__load_fee_payments(person_id)
        self.__load_fine_payments(person_id)
        self.__update_payment_summary(person_id)

    def reject(self):

        self.rollback()
        QDialog.reject(self)

    def __load_fee_payments(self, person_id):

        self.payment_twidget.setRowCount(0)
        count = self.contract.fees.filter(CtFee.person == person_id).count()
        if count == 0:
            return

        payment_year = self.year_sbox.value()
        fee = self.contract.fees.filter(CtFee.person == person_id).one()
        count = fee.payments.filter(CtFeePayment.year_paid_for == payment_year).count()
        self.payment_twidget.setRowCount(count)
        row = 0

        for payment in fee.payments.filter(CtFeePayment.year_paid_for == payment_year).order_by(CtFeePayment.date_paid):

            self.__add_payment_row(row, payment, self.payment_twidget)
            row += 1

        if row > 0:
            self.payment_twidget.resizeColumnToContents(PAYMENT_DATE)
            self.payment_twidget.resizeColumnToContents(AMOUNT_PAID)
            self.payment_twidget.horizontalHeader().setResizeMode(PAYMENT_TYPE, QHeaderView.Stretch)

    def __load_fine_payments(self, person_id):

        self.fine_payment_twidget.setRowCount(0)
        count = self.contract.fees.filter(CtFee.person == person_id).count()
        if count == 0:
            return

        payment_year = self.year_sbox.value()
        fee = self.contract.fees.filter(CtFee.person == person_id).one()
        count = fee.fine_payments.filter(CtFineForFeePayment.year_paid_for == payment_year).count()
        self.fine_payment_twidget.setRowCount(count)
        row = 0

        for payment in fee.fine_payments.filter(CtFineForFeePayment.year_paid_for == payment_year)\
                .order_by(CtFineForFeePayment.date_paid):

            self.__add_payment_row(row, payment, self.fine_payment_twidget)
            row += 1

        if row > 0:
            self.fine_payment_twidget.resizeColumnToContents(PAYMENT_DATE)
            self.fine_payment_twidget.resizeColumnToContents(AMOUNT_PAID)
            self.fine_payment_twidget.horizontalHeader().setResizeMode(PAYMENT_TYPE, QHeaderView.Stretch)

    def __add_payment_row(self, row, payment, twidget):

        item = QTableWidgetItem('{0}'.format(PluginUtils.convert_python_date_to_qt(payment.date_paid)
                                             .toString(Constants.DATE_FORMAT)))
        item.setData(Qt.UserRole, payment.id)
        self.__lock_item(item)
        twidget.setItem(row, PAYMENT_DATE, item)

        item = QTableWidgetItem('{0}'.format(payment.amount_paid))
        self.__lock_item(item)
        twidget.setItem(row, AMOUNT_PAID, item)

        item = QTableWidgetItem(u'{0}'.format(payment.payment_type_ref.description))
        item.setData(Qt.UserRole, payment.payment_type_ref.code)
        self.__lock_item(item)
        twidget.setItem(row, PAYMENT_TYPE, item)

    def __lock_item(self, item):

        item.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)

    @pyqtSlot()
    def on_register_payment_button_clicked(self):

        # Validate
        person_id = self.select_contractor_cbox.itemData(self.select_contractor_cbox.currentIndex(), Qt.UserRole)
        count = self.contract.fees.filter(CtFee.person == person_id).count()
        if count == 0:
            PluginUtils.show_message(self, self.tr("Register Payment"),
                                     self.tr("No fee is registered for the selected contractor!"))
            return

        if self.contract.status != 20:
            PluginUtils.show_message(self, self.tr("Register Payment"),
                                     self.tr("Payments can be registered for active contracts only!"))
            return

        contract_begin_year = self.contract.contract_begin.year
        contract_end_year = self.contract.contract_end.year

        effective_fine = int(self.effective_fine_edit.text())
        if effective_fine > 0:
            PluginUtils.show_message(self, self.tr("Register Payment"),
                                     self.tr("No more payments can be registered for this payment year!"))
            return

        payment_year = self.year_sbox.value()
        if payment_year < contract_begin_year or payment_year > contract_end_year:
            PluginUtils.show_message(self, self.tr("Register Payment"),
                                     self.tr("Payments cannot be registered for years outside the contract period!"))
            return

        fee = self.contract.fees.filter(CtFee.person == person_id).one()
        payment_date = self.payment_date_edit.date()
        amount_paid = self.amount_paid_sbox.value()
        payment_type = self.payment_type_cbox.itemData(self.payment_type_cbox.currentIndex(), Qt.UserRole)

        if amount_paid == 0:
            PluginUtils.show_message(self, self.tr("Register Payment"),
                                     self.tr("The amount paid must be greater than 0!"))
            return

        payment = CtFeePayment()
        payment.date_paid = PluginUtils.convert_qt_date_to_python(payment_date)
        payment.amount_paid = amount_paid
        payment.payment_type = payment_type
        payment.paid_total = 0
        payment.year_paid_for = payment_year

        fee.payments.append(payment)
        self.session.flush()
        self.__load_fee_payments(person_id)
        self.__calculate_fines(fee, payment_year)
        self.__update_payment_summary(person_id)

    @pyqtSlot()
    def on_remove_button_clicked(self):

        selected_row = self.payment_twidget.currentRow()
        if selected_row == -1:
            PluginUtils.show_message(self, self.tr("Select Payment"),
                                     self.tr("Please select payment!"))
            return
        date_item = self.payment_twidget.item(selected_row, 0)
        payment_id = date_item.data(Qt.UserRole)
        amount_item = self.payment_twidget.item(selected_row, 1)
        type_item = self.payment_twidget.item(selected_row, 2)
        self.payment_twidget.removeRow(selected_row)

        person_id = self.select_contractor_cbox.itemData(self.select_contractor_cbox.currentIndex(), Qt.UserRole)
        count = self.contract.fees.filter(CtFee.person == person_id).count()
        if count == 0:
            PluginUtils.show_message(self, self.tr("Register Payment"),
                                     self.tr("No fee is registered for the selected contractor!"))
            return

        fee = self.contract.fees.filter(CtFee.person == person_id).one()
        payment_year = self.year_sbox.value()
        payment = self.session.query(CtFeePayment).filter(CtFeePayment.id == payment_id).one()

        fee.payments.remove(payment)
        self.session.flush()
        self.__load_fee_payments(person_id)
        self.__calculate_fines(fee, payment_year)
        self.__update_payment_summary(person_id)

    @pyqtSlot(QTableWidgetItem)
    def on_payment_twidget_itemClicked(self, item):

        current_row = self.payment_twidget.currentRow()

        date_item = self.payment_twidget.item(current_row, 0)
        payment_id = date_item.data(Qt.UserRole)
        payment = self.session.query(CtFeePayment).filter(CtFeePayment.id == payment_id).one()
        qt_date = PluginUtils.convert_python_date_to_qt(payment.date_paid)

        amount_item = self.payment_twidget.item(current_row, 1)
        type_item = self.payment_twidget.item(current_row, 2)
        payment_type = type_item.data(Qt.UserRole)

        self.payment_date_edit.setDate(qt_date)
        self.amount_paid_sbox.setValue(int(amount_item.text()))
        self.payment_type_cbox.setCurrentIndex(self.payment_type_cbox.findData(payment_type))

    @pyqtSlot()
    def on_edit_button_clicked(self):

        person_id = self.select_contractor_cbox.itemData(self.select_contractor_cbox.currentIndex(), Qt.UserRole)
        count = self.contract.fees.filter(CtFee.person == person_id).count()
        if count == 0:
            PluginUtils.show_message(self, self.tr("Register Payment"),
                                     self.tr("No fee is registered for the selected contractor!"))
            return

        fee = self.contract.fees.filter(CtFee.person == person_id).one()

        current_row = self.payment_twidget.currentRow()

        date_item = self.payment_twidget.item(current_row, 0)
        if date_item == None:
            PluginUtils.show_message(self, self.tr("Select Payment"),
                                     self.tr("Please select payment!"))
            return
        payment_id = date_item.data(Qt.UserRole)

        payment_date = self.payment_date_edit.date()
        amount_paid = self.amount_paid_sbox.value()
        payment_type = self.payment_type_cbox.itemData(self.payment_type_cbox.currentIndex(), Qt.UserRole)
        payment_year = self.year_sbox.value()

        payment = self.session.query(CtFeePayment).filter(CtFeePayment.id == payment_id).one()
        payment.date_paid = PluginUtils.convert_qt_date_to_python(payment_date)
        payment.amount_paid = amount_paid
        payment.payment_type = payment_type
        payment.paid_total = 0
        payment.year_paid_for = payment_year

        # fee.payments.extend(payment)
        self.session.flush()

        date_item = self.payment_twidget.item(current_row, 0)
        amount_item = self.payment_twidget.item(current_row, 1)
        type_item = self.payment_twidget.item(current_row, 2)

        date_item.setText(str(self.payment_date_edit.text()))
        date_item.setData(Qt.UserRole,payment_id)
        amount_item.setText(str(self.amount_paid_sbox.value()))
        amount_item.setData(Qt.UserRole, self.amount_paid_sbox.value())
        type_item.setText(self.payment_type_cbox.currentText())
        type_item.setData(Qt.UserRole, payment_type)

        self.__load_fee_payments(person_id)
        self.__calculate_fines(fee, payment_year)
        self.__update_payment_summary(person_id)

    def __calculate_fines(self, fee, payment_year):

        fine_rate = self.session.query(SetPayment.landfee_fine_rate_per_day).first()[0]
        surplus_from_previous_years = self.__surplus_from_previous_years(fee)

        payments = fee.payments.filter(CtFeePayment.year_paid_for == payment_year)\
            .order_by(CtFeePayment.date_paid)

        paid_total = surplus_from_previous_years

        grace_period = int(self.grace_period_edit.text())

        for payment in payments:

            delay_to_dl1 = QDate(payment_year, 1, 25).addDays(grace_period)\
                .daysTo(QDate(payment.date_paid.year, payment.date_paid.month, payment.date_paid.day))
            if delay_to_dl1 > 0:
                payment.delay_to_dl1 = delay_to_dl1
            else:
                delay_to_dl1 = 0
            delay_to_dl2 = QDate(payment_year, 4, 25).addDays(grace_period)\
                .daysTo(QDate(payment.date_paid.year, payment.date_paid.month, payment.date_paid.day))
            if delay_to_dl2 > 0:
                payment.delay_to_dl2 = delay_to_dl2
            else:
                delay_to_dl2 = 0
            delay_to_dl3 = QDate(payment_year, 7, 25).addDays(grace_period)\
                .daysTo(QDate(payment.date_paid.year, payment.date_paid.month, payment.date_paid.day))
            if delay_to_dl3 > 0:
                payment.delay_to_dl3 = delay_to_dl3
            else:
                delay_to_dl3 = 0
            delay_to_dl4 = QDate(payment_year, 10, 25).addDays(grace_period)\
                .daysTo(QDate(payment.date_paid.year, payment.date_paid.month, payment.date_paid.day))
            if delay_to_dl4 > 0:
                payment.delay_to_dl4 = delay_to_dl4
            else:
                delay_to_dl4 = 0

            left_to_pay = self.__left_to_pay(fee, paid_total)

            fine_for_q1 = delay_to_dl1 * left_to_pay['Q1'] * fine_rate / 100
            fine_for_q2 = delay_to_dl2 * left_to_pay['Q2'] * fine_rate / 100
            fine_for_q3 = delay_to_dl3 * left_to_pay['Q3'] * fine_rate / 100
            fine_for_q4 = delay_to_dl4 * left_to_pay['Q4'] * fine_rate / 100

            if fine_for_q1 > left_to_pay['Q1'] / 2:
                fine_for_q1 = left_to_pay['Q1'] / 2
            if fine_for_q2 > left_to_pay['Q2'] / 2:
                fine_for_q2 = left_to_pay['Q2'] / 2
            if fine_for_q3 > left_to_pay['Q3'] / 2:
                fine_for_q3 = left_to_pay['Q3'] / 2
            if fine_for_q4 > left_to_pay['Q4'] / 2:
                fine_for_q4 = left_to_pay['Q4'] / 2

            payment.fine_for_q1 = fine_for_q1
            payment.fine_for_q2 = fine_for_q2
            payment.fine_for_q3 = fine_for_q3
            payment.fine_for_q4 = fine_for_q4

            payment.paid_total = payment.amount_paid + paid_total
            paid_total += payment.amount_paid

            left_to_pay = self.__left_to_pay(fee, paid_total)

            payment.left_to_pay_for_q1 = left_to_pay['Q1']
            payment.left_to_pay_for_q2 = left_to_pay['Q2']
            payment.left_to_pay_for_q3 = left_to_pay['Q3']
            payment.left_to_pay_for_q4 = left_to_pay['Q4']

        self.session.flush()

    def __left_to_pay(self, fee, paid_total):

        payment_year = self.year_sbox.value()

        fee_to_pay_for_q1 = self.__fee_to_pay_per_period(fee, date(payment_year, 1, 1), date(payment_year, 4, 1))
        fee_to_pay_for_q2 = self.__fee_to_pay_per_period(fee, date(payment_year, 4, 1), date(payment_year, 7, 1))
        fee_to_pay_for_q3 = self.__fee_to_pay_per_period(fee, date(payment_year, 7, 1), date(payment_year, 10, 1))
        fee_to_pay_for_q4 = self.__fee_to_pay_per_period(fee, date(payment_year, 10, 1), date(payment_year+1, 1, 1))

        left_to_pay_for_q1 = fee_to_pay_for_q1 - paid_total
        if left_to_pay_for_q1 < 0:
            left_to_pay_for_q1 = 0
        surplus = paid_total - fee_to_pay_for_q1
        if surplus < 0:
            surplus = 0
        left_to_pay_for_q2 = fee_to_pay_for_q2 - surplus
        if left_to_pay_for_q2 < 0:
            left_to_pay_for_q2 = 0
        surplus = paid_total - (fee_to_pay_for_q1 + fee_to_pay_for_q2)
        if surplus < 0:
            surplus = 0
        left_to_pay_for_q3 = fee_to_pay_for_q3 - surplus
        if left_to_pay_for_q3 < 0:
            left_to_pay_for_q3 = 0
        surplus = paid_total - (fee_to_pay_for_q1 + fee_to_pay_for_q2 + fee_to_pay_for_q3)
        if surplus < 0:
            surplus = 0
        left_to_pay_for_q4 = fee_to_pay_for_q4 - surplus
        if left_to_pay_for_q4 < 0:
            left_to_pay_for_q4 = 0

        left_to_pay = dict()
        left_to_pay['Q1'] = left_to_pay_for_q1
        left_to_pay['Q2'] = left_to_pay_for_q2
        left_to_pay['Q3'] = left_to_pay_for_q3
        left_to_pay['Q4'] = left_to_pay_for_q4

        return left_to_pay

    @pyqtSlot()
    def on_fine_register_payment_button_clicked(self):

        # Validate
        person_id = self.select_contractor_cbox.itemData(self.select_contractor_cbox.currentIndex(), Qt.UserRole)
        count = self.contract.fees.filter(CtFee.person == person_id).count()
        if count == 0:
            PluginUtils.show_message(self, self.tr("Register Fine Payment"),
                                     self.tr("No fee is registered for the selected contractor!"))
            return

        if self.contract.status != 20:
            PluginUtils.show_message(self, self.tr("Register Fine Payment"),
                                     self.tr("Payments can be registered for active contracts only!"))
            return

        contract_begin_year = self.contract.contract_begin.year
        contract_end_year = self.contract.contract_end.year

        payment_year = self.year_sbox.value()
        if payment_year < contract_begin_year or payment_year > contract_end_year:
            PluginUtils.show_message(self, self.tr("Register Fine Payment"),
                                     self.tr("Payments cannot be registered for years outside the contract period!"))
            return

        effective_fine = int(self.effective_fine_edit.text())
        if effective_fine == 0:
            PluginUtils.show_message(self, self.tr("Register Fine Payment"),
                                     self.tr("A fine payment cannot be registered without an effective fine!"))
            return

        fee = self.contract.fees.filter(CtFee.person == person_id).one()
        payment_date = self.fine_payment_date_edit.date()
        amount_paid = self.fine_amount_paid_sbox.value()
        payment_type = self.fine_payment_type_cbox.itemData(self.payment_type_cbox.currentIndex(), Qt.UserRole)

        if amount_paid == 0:
            PluginUtils.show_message(self, self.tr("Register Fine Payment"),
                                     self.tr("The amount paid must be greater than 0!"))
            return

        payment = CtFineForFeePayment()
        payment.date_paid = PluginUtils.convert_qt_date_to_python(payment_date)
        payment.amount_paid = amount_paid
        payment.payment_type = payment_type
        payment.year_paid_for = payment_year

        fee.fine_payments.append(payment)
        self.session.flush()
        self.__load_fine_payments(person_id)
        self.__update_payment_summary(person_id)

    def __update_payment_summary(self, person_id):

        count = self.contract.fees.filter(CtFee.person == person_id).count()
        if count == 0:
            return

        fee = self.contract.fees.filter(CtFee.person == person_id).one()
        self.grace_period_edit.setText(str(fee.grace_period))
        self.payment_frequency_edit.setText(fee.payment_frequency_ref.description)

        self.__set_fee_summary(fee)
        self.__set_fine_summary(fee)
        self.__update_payment_status(fee)

    def __fee_to_pay_per_period(self, fee, period_begin, period_end):

        # Intersect contract duration with payment period
        sql = "select lower(daterange(contract_begin, contract_end, '[)') * daterange(:from, :to, '[)'))," \
              " upper(daterange(contract_begin, contract_end, '[)') * daterange(:from, :to, '[)')) " \
              "from ct_contract where contract_id = :contract_id"

        result = self.session.execute(sql, {'from': period_begin,
                                            'to': period_end,
                                            'contract_id': fee.contract})
        for row in result:
            effective_begin = row[0]
            effective_end = row[1]

        if effective_begin is None and effective_end is None:
            return 0

        # Intersect the effective payment period with the archived fees
        sql = "select upper(daterange(valid_from, valid_till, '[)') * daterange(:begin, :end, '[)')) - "\
                 "lower(daterange(valid_from, valid_till, '[)') * daterange(:begin, :end, '[)')) as days, "\
                 "fee_contract from ct_archived_fee where contract = :contract and person = :person"

        result = self.session.execute(sql, {'begin': effective_begin,
                                            'end': effective_end,
                                            'contract': fee.contract,
                                            'person': fee.person})
        fee_for_period = 0
        total_days = 0

        for row in result:
            days = row[0]
            if days is None:
                continue
            annual_fee = row[1]
            adjusted_fee = (annual_fee / 365) * days
            fee_for_period += adjusted_fee
            total_days += days

        effective_days = (effective_end-effective_begin).days

        if effective_days - total_days > 0:
            fee_for_period += (effective_days-total_days) * fee.fee_contract / 365

        return int(round(fee_for_period))

    def __set_fee_summary(self, fee):

        payment_year = self.year_sbox.value()

        fee_to_pay_for_current_year = \
            self.__fee_to_pay_per_period(fee, date(payment_year, 1, 1), date(payment_year+1, 1, 1))

        paid_for_current_year = self.session.query(func.sum(CtFeePayment.amount_paid))\
            .filter(CtFeePayment.contract == fee.contract).filter(CtFeePayment.person == fee.person)\
            .filter(CtFeePayment.year_paid_for == payment_year).scalar()

        if paid_for_current_year is None:
            paid_for_current_year = 0

        surplus = self.__surplus_from_previous_years(fee)

        fee_left_to_pay = fee_to_pay_for_current_year - (paid_for_current_year + surplus)
        if fee_left_to_pay < 0:
            fee_left_to_pay = 0

        # set for display
        self.fee_per_year_edit.setText(str(fee_to_pay_for_current_year))
        self.fee_paid_edit.setText(str(paid_for_current_year))
        self.surplus_from_last_years_edit.setText(str(surplus))
        self.fee_to_pay_edit.setText(str(fee_left_to_pay))
        if fee_left_to_pay > 0:
            self.__change_text_color(self.fee_to_pay_edit)
        else:
            self.__reset_text_color(self.fee_to_pay_edit)

    def __set_fine_summary(self, fee):

        payment_year = self.year_sbox.value()

        effective_fine_for_current_year = self.__effective_fine_for_year(fee, payment_year)
        potential_fine_for_current_year = self.__potential_fine_for_year(fee, payment_year)

        paid_for_current_year = self.session.query(func.sum(CtFineForFeePayment.amount_paid))\
            .filter(CtFineForFeePayment.contract == fee.contract).filter(CtFineForFeePayment.person == fee.person)\
            .filter(CtFineForFeePayment.year_paid_for == payment_year).scalar()

        if paid_for_current_year is None:
            paid_for_current_year = 0

        self.effective_fine_edit.setText(str(effective_fine_for_current_year))
        self.potential_fine_edit.setText(str(potential_fine_for_current_year))
        self.fine_paid_edit.setText(str(paid_for_current_year))
        fine_to_pay = effective_fine_for_current_year - paid_for_current_year
        if fine_to_pay < 0:
            fine_to_pay = 0
        self.fine_to_pay_edit.setText(str(fine_to_pay))
        if fine_to_pay > 0:
            self.__change_text_color(self.fine_to_pay_edit)
        else:
            self.__reset_text_color(self.fine_to_pay_edit)

    def __effective_fine_for_year(self, fee, payment_year):

        return self.__total_fine(fee, payment_year)

    def __potential_fine_for_year(self, fee, payment_year):

        return self.__total_fine(fee, payment_year, False)

    def __total_fine(self, fee, payment_year, effective_fine=True):

        count = self.session.query(CtFeePayment)\
            .filter(CtFeePayment.contract == fee.contract)\
            .filter(CtFeePayment.person == fee.person)\
            .filter(CtFeePayment.year_paid_for == payment_year)\
            .filter(CtFeePayment.left_to_pay_for_q1 == 0)\
            .filter(CtFeePayment.left_to_pay_for_q2 == 0)\
            .filter(CtFeePayment.left_to_pay_for_q3 == 0)\
            .filter(CtFeePayment.left_to_pay_for_q4 == 0).count()

        if effective_fine:
            if count == 0:
                return 0
        else:
            if count != 0:
                return 0

        payment_frequency = fee.payment_frequency
        total_fine = 0
        fine = self.session.query(func.sum(CtFeePayment.fine_for_q1))\
            .filter(CtFeePayment.contract == fee.contract).filter(CtFeePayment.person == fee.person)\
            .filter(CtFeePayment.year_paid_for == payment_year).scalar()
        if fine is not None and payment_frequency == 20:
            total_fine += fine
        fine = self.session.query(func.sum(CtFeePayment.fine_for_q2))\
            .filter(CtFeePayment.contract == fee.contract).filter(CtFeePayment.person == fee.person)\
            .filter(CtFeePayment.year_paid_for == payment_year).scalar()
        if fine is not None and payment_frequency == 20:
            total_fine += fine
        fine = self.session.query(func.sum(CtFeePayment.fine_for_q3))\
            .filter(CtFeePayment.contract == fee.contract).filter(CtFeePayment.person == fee.person)\
            .filter(CtFeePayment.year_paid_for == payment_year).scalar()
        if fine is not None and payment_frequency == 20:
            total_fine += fine
        fine = self.session.query(func.sum(CtFeePayment.fine_for_q4))\
            .filter(CtFeePayment.contract == fee.contract).filter(CtFeePayment.person == fee.person)\
            .filter(CtFeePayment.year_paid_for == payment_year).scalar()
        if fine is not None:
            total_fine += fine

        return int(round(total_fine))

    def __surplus_from_previous_years(self, fee):

        year_to_pay_for = self.year_sbox.value()

        surplus_last_year = 0

        for payment_year in range(self.contract.contract_begin.year, year_to_pay_for):

            amount_paid = self.session.query(func.sum(CtFeePayment.amount_paid))\
                .filter(CtFeePayment.contract == fee.contract).filter(CtFeePayment.person == fee.person)\
                .filter(CtFeePayment.year_paid_for == payment_year).scalar()
            if amount_paid is None:
                amount_paid = 0

            fee_to_pay = self.__fee_to_pay_per_period(fee, date(payment_year, 1, 1), date(payment_year+1, 1, 1))
            if (amount_paid + surplus_last_year) - fee_to_pay > 0:
                surplus_last_year = (amount_paid + surplus_last_year) - fee_to_pay
            else:
                surplus_last_year = 0

        return surplus_last_year

    @pyqtSlot()
    def on_apply_button_clicked(self):

        self.commit()
        self.__start_fade_out_timer()

    def __start_fade_out_timer(self):

        self.timer = QTimer()
        self.timer.timeout.connect(self.__fade_status_message)
        self.time_counter = 500
        self.timer.start(10)

    def __fade_status_message(self):

        opacity = int(self.time_counter * 0.5)
        self.status_label.setStyleSheet("QLabel {color: rgba(255,0,0," + str(opacity) + ");}")
        self.status_label.setText(self.tr('Changes applied successfully.'))
        if self.time_counter == 0:
            self.timer.stop()
        self.time_counter -= 1

    @pyqtSlot(int)
    def on_year_sbox_valueChanged(self, sbox_value):

        self.__clear_controls()
        person_id = self.select_contractor_cbox.itemData(self.select_contractor_cbox.currentIndex(), Qt.UserRole)
        count = self.contract.fees.filter(CtFee.person == person_id).count()
        if count == 0:
            return

        self.__load_fee_payments(person_id)
        self.__load_fine_payments(person_id)
        self.__update_payment_summary(person_id)

    def __xxx_fine(self, fee):

        fee_left_to_be_paid = int(self.fee_to_pay_edit.text())
        if fee_left_to_be_paid == 0:
            return 0

        amount_paid = int(self.fee_paid_edit.text()) + int(self.surplus_from_last_year_edit.text())

        payment_year = self.year_sbox.value()

        payment_date = self.payment_date_edit.date()

        grace_period = int(self.grace_period_edit.text())

        fee_to_pay_for_q1 = self.__fee_to_pay_per_period(fee, date(payment_year, 1, 1), date(payment_year, 4, 1))
        fee_to_pay_for_q2 = self.__fee_to_pay_per_period(fee, date(payment_year, 4, 1), date(payment_year, 7, 1))
        fee_to_pay_for_q3 = self.__fee_to_pay_per_period(fee, date(payment_year, 7, 1), date(payment_year, 10, 1))
        fee_to_pay_for_q4 = self.__fee_to_pay_per_period(fee, date(payment_year, 10, 1), date(payment_year+1, 1, 1))

        if fee.payment_frequency == 20:  # Quarterly payment

            if fee_to_pay_for_q1 > amount_paid:
                deadline = QDate(payment_year, 1, 25)
            elif fee_to_pay_for_q1+fee_to_pay_for_q2 > amount_paid:
                deadline = QDate(payment_year, 4, 25)
            elif fee_to_pay_for_q1+fee_to_pay_for_q2+fee_to_pay_for_q3 > amount_paid:
                deadline = QDate(payment_year, 7, 25)
            elif fee_to_pay_for_q1+fee_to_pay_for_q2+fee_to_pay_for_q3+fee_to_pay_for_q4 > amount_paid:
                deadline = QDate(payment_year, 10, 25)
        else:  # Annual payment

            deadline = QDate(payment_year, 10, 25)

        delayed_days = deadline.addDays(grace_period).daysTo(payment_date)

        if delayed_days < 0:
            delayed_days = 0

        fine_rate = self.session.query(SetPayment.landfee_fine_rate_per_day).first()[0]

        fine = delayed_days * (fine_rate / 100) * fee_left_to_be_paid

        if fine > fee_left_to_be_paid / 2:
            fine = fee_left_to_be_paid / 2

        fine = int(round(fine))

        #self.delayed_payment_edit.setText(str(delayed_days))

        return fine

    def __update_payment_status(self, fee):

        amount_paid = int(self.fee_paid_edit.text()) + int(self.surplus_from_last_years_edit.text())

        payment_year = self.year_sbox.value()

        fee_to_pay_for_q1 = self.__fee_to_pay_per_period(fee, date(payment_year, 1, 1), date(payment_year, 4, 1))
        fee_to_pay_for_q2 = self.__fee_to_pay_per_period(fee, date(payment_year, 4, 1), date(payment_year, 7, 1))
        fee_to_pay_for_q3 = self.__fee_to_pay_per_period(fee, date(payment_year, 7, 1), date(payment_year, 10, 1))
        fee_to_pay_for_q4 = self.__fee_to_pay_per_period(fee, date(payment_year, 10, 1), date(payment_year+1, 1, 1))

        if 0 < fee_to_pay_for_q1 <= amount_paid:
            self.quarter_1_check_box.setChecked(True)

        if fee_to_pay_for_q2 > 0 and amount_paid >= fee_to_pay_for_q1+fee_to_pay_for_q2:
            self.quarter_2_check_box.setChecked(True)

        if fee_to_pay_for_q3 > 0 and amount_paid >= fee_to_pay_for_q1+fee_to_pay_for_q2+fee_to_pay_for_q3:
            self.quarter_3_check_box.setChecked(True)

        if fee_to_pay_for_q4 > 0 and \
                amount_paid >= fee_to_pay_for_q1+fee_to_pay_for_q2+fee_to_pay_for_q3+fee_to_pay_for_q4:
            self.quarter_4_check_box.setChecked(True)

    def __clear_controls(self):

        self.grace_period_edit.setText('0')
        self.payment_frequency_edit.setText('0')
        self.fee_per_year_edit.setText('0')
        self.fee_paid_edit.setText('0')
        self.surplus_from_last_years_edit.setText('0')
        self.fee_to_pay_edit.setText('0')
        self.potential_fine_edit.setText('0')
        self.effective_fine_edit.setText('0')
        self.fine_paid_edit.setText('0')
        self.fine_to_pay_edit.setText('0')
        self.payment_twidget.setRowCount(0)
        self.fine_payment_twidget.setRowCount(0)
        self.quarter_1_check_box.setChecked(False)
        self.quarter_2_check_box.setChecked(False)
        self.quarter_3_check_box.setChecked(False)
        self.quarter_4_check_box.setChecked(False)

        self.__reset_text_color(self.fee_to_pay_edit)
        self.__reset_text_color(self.fine_to_pay_edit)

    def __change_text_color(self, line_edit):

        style_sheet = "QLineEdit {color:rgb(255, 0, 0);}"
        line_edit.setStyleSheet(style_sheet)

    def __reset_text_color(self, line_edit):

        line_edit.setStyleSheet(None)

    @pyqtSlot()
    def on_help_button_clicked(self):

        if self.tabWidget.currentIndex() == 0:
                os.system("hh.exe "+ str(os.path.dirname(os.path.realpath(__file__))[:-10]) +"help\output\help_lm2.chm::/html/summary.htm")
        elif self.tabWidget.currentIndex() == 1:
                os.system("hh.exe "+ str(os.path.dirname(os.path.realpath(__file__))[:-10]) +"help\output\help_lm2.chm::/html/fees.htm")
        elif self.tabWidget.currentIndex() == 2:
                os.system("hh.exe "+ str(os.path.dirname(os.path.realpath(__file__))[:-10]) +"help\output\help_lm2.chm::/html/fines.htm")

    def set_column_width(self, column, width):
        for cell in column:
            cell.width = width

    # table = document.add_table(rows=10,cols=2,style="ColorfulList-Accent5")

    def __create_fee_unifeid_view(self):

        current_working_soum = "'" + str(DatabaseUtils.current_working_soum_schema()) + "'"

        sql = ""

        if not sql:
            sql = "Create temp view land_fee_unified as" + "\n"
        else:
            sql = sql + "UNION" + "\n"

        select = "SELECT contract.contract_id, contract.contract_no, payment.year_paid_for,contract.status, fee.fee_contract, sum(payment.amount_paid) as paid, person.person_id, p_paid.p_paid ,landuse.description as landuse, " \
                    "decision.decision_date ,decision.decision_no, person.first_name, person.name, person.contact_surname, person.contact_first_name ,person.address_street_name as person_streetname, person.address_khaskhaa as person_khashaa, " \
                    "parcel.parcel_id, contract.certificate_no, person.person_register ,au3_person.name as person_bag,au3.name as bag_name, person.mobile_phone, parcel.area_m2, application.approved_duration, parcel.address_streetname, parcel.address_khashaa " \
                 "FROM data_soums_union.ct_contract contract " \
                 "LEFT JOIN data_soums_union.ct_contract_application_role con_app on con_app.contract = contract.contract_id "\
                 "LEFT JOIN data_soums_union.ct_application application ON application.app_id = con_app.application " \
                 "LEFT JOIN data_soums_union.ct_application_person_role app_pers on application.app_id = app_pers.application "\
                 "LEFT JOIN base.bs_person person ON app_pers.person = person.person_id " \
                 "LEFT JOIN data_soums_union.ca_parcel_tbl parcel on parcel.parcel_id = application.parcel " \
                 "LEFT JOIN admin_units.au_level3 au3 on ST_Within(parcel.geometry,au3.geometry) "\
                 "LEFT JOIN admin_units.au_level3 au3_person on person.address_au_level3 = au3_person.code " \
                 "LEFT JOIN codelists.cl_landuse_type landuse on parcel.landuse = landuse.code " \
                 "LEFT JOIN data_soums_union.ct_fee fee on contract.contract_id = fee.contract " \
                 "LEFT JOIN data_soums_union.ct_decision_application dec_app on dec_app.application = application.app_id " \
                 "LEFT JOIN data_soums_union.ct_decision decision on decision.decision_id = dec_app.decision " \
                 "LEFT JOIN (select p.contract, sum(p.amount_paid) as p_paid from data_soums_union.ct_fee_payment p where p.year_paid_for < date_part('year', NOW())::int group by p.contract) p_paid on contract.contract_id = p_paid.contract " \
                 "LEFT JOIN data_soums_union.ct_fee_payment payment on fee.contract = payment.contract " \
                 "where  application.au2 = {0}".format(current_working_soum) + "\n"

        sql = sql + select
        sql = "{0} group by person.person_register, contract.contract_id, contract.contract_no, payment.year_paid_for,contract.status,fee.fee_contract, person.person_id, p_paid.p_paid,decision.decision_no, decision.decision_date, contract.certificate_no, parcel.parcel_id,au3.name, application.approved_duration,landuse.description,parcel.area_m2, parcel.address_streetname, parcel.address_khashaa,au3_person.name ".format(sql)
        sql = "{0}  order by contract_no;".format(sql)
        # try:
        self.session.execute(sql)
        self.commit()

        # except SQLAlchemyError, e:
        #     PluginUtils.show_message(self, self.tr("LM2", "Sql Error"), e.message)
        #     return

    @pyqtSlot()
    def on_act_print_button_clicked(self):

        self.__create_fee_unifeid_view()
        person_id = self.select_contractor_cbox.itemData(self.select_contractor_cbox.currentIndex(), Qt.UserRole)
        if not person_id:
            PluginUtils.show_message(self, self.tr("Land Fee"),
                                     self.tr(" Land fee is not registered!"))
            return
        fee = self.contract.fees.filter(CtFee.person == person_id).one()
        self.__act_surplus_from_previous_years(fee)

    def __set_column_width(self, column, width):
        for cell in column.cells:
            cell.width = width

    def __act_surplus_from_previous_years(self, fee):

        document = Document()
        sections = document.sections
        sections[0].left_margin = Cm(1.5)
        sections[0].right_margin = Cm(1.5)
        sections[0].top_margin = Cm(1.5)
        sections[0].bottom_margin = Cm(1.5)

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


        p_h1 = p.add_run(u'ГАЗРЫН ТӨРБӨРИЙН ТООЦОО НИЙЛСЭН АКТ  \n')
        font = p_h1.font
        font.name = 'Calibri'
        font.size = Pt(12)
        font.bold = True

        p_h = p.add_run(u'№ ')
        font = p_h.font
        font.name = 'Calibri'
        font.size = Pt(12)
        font.bold = True

        info = self.session.query(FeeUnified). \
            filter(FeeUnified.contract_no == self.contract.contract_no). \
            filter(FeeUnified.person_id == fee.person).one()

        table = document.add_table(rows=11, cols=2)
        # table.style = 'TableNormal'
        # table.style = 'TableGrid'
        table.style = 'LightShading-Accent1'
        # table.allow_autofit = True
        # table.autofit = True

        a = table.rows[0].cells[0]
        b = table.rows[0].cells[1]
        A = a.merge(b)
        cell = table.rows[0].cells[0]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(u'Эрх зүйн харилцааны мэдээлэл')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True
        ############
        cell = table.rows[1].cells[0]
        cell.width = Cm(0.5)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(u'Засаг даргын шийдвэрийн огноо:')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True

        cell = table.rows[1].cells[1]
        cell.width = Cm(0.5)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(str(info.decision_date) +', ' + info.decision_no)
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True
        #######################
        cell = table.rows[2].cells[0]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(u'Газар эзэмшигч иргэн, хуулийн этгээдийн нэр:')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True

        cell = table.rows[2].cells[1]
        cell.width = Cm(0.5)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name = ''
        if info.name:
            name = info.name
        first_name = ''
        if info.first_name:
            first_name = info.first_name
        run = paragraph.add_run(name +' '+ first_name)
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True
        ###########################
        cell = table.rows[3].cells[0]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(u'Регистрийн дугаар:')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True

        cell = table.rows[3].cells[1]
        cell.width = Cm(0.5)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(info.person_register)
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True
        #########################
        cell = table.rows[4].cells[0]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(u'Гэрийн хаяг:')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True

        cell = table.rows[4].cells[1]
        cell.width = Cm(0.5)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        person_bag = ''
        if info.person_bag:
            person_bag = info.person_bag
        person_streetname = ''
        if info.person_streetname:
            person_streetname = info.person_streetname
        person_khashaa = ''
        if info.person_khashaa:
            person_khashaa = info.person_khashaa
        run = paragraph.add_run(person_bag +' '+ person_streetname +' '+ person_khashaa)
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True
        #########################
        cell = table.rows[5].cells[0]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(u'Нэгж талбарын дугаар, гэрчилгээний дугаар:')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True

        cell = table.rows[5].cells[1]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(info.parcel_id +' '+ str(info.certificate_no))
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True
        ########################
        cell = table.rows[6].cells[0]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(u'Газрын байршил, зориулалт:')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True

        cell = table.rows[6].cells[1]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        address_streetname = ''
        if info.address_streetname:
            address_streetname = info.address_streetname
        address_khashaa = ''
        if info.address_khashaa:
            address_khashaa = info.address_khashaa
        landuse = ''
        if info.landuse:
            landuse = info.landuse
        bag_name = ''
        if info.bag_name:
            bag_name = info.bag_name
        run = paragraph.add_run(bag_name +' '+ address_streetname +' '+ address_khashaa +', '+ landuse)
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True
        ##########################
        cell = table.rows[7].cells[0]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(u'Утасны дугаар:')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True

        cell = table.rows[7].cells[1]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        mobile_phone = ''
        if info.mobile_phone:
            mobile_phone = info.mobile_phone
        run = paragraph.add_run(mobile_phone)
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True
        ###########################
        cell = table.rows[8].cells[0]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(u'Талбайн хэмжээ:')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True

        cell = table.rows[8].cells[1]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(str(info.area_m2))
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True
        ##########################
        cell = table.rows[9].cells[0]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(u'Газрын жилийн төлбөр:')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True

        cell = table.rows[9].cells[1]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(str(info.fee_contract))
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True
        ##########################
        cell = table.rows[10].cells[0]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(u'Газар эзэмшин хугацаа:')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True

        cell = table.rows[10].cells[1]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(str(info.approved_duration))
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True

        p = document.add_paragraph()
        date_now = QDate.currentDate()
        date_now = date_now.toString(Constants.DATE_FORMAT)
        run5 = p.add_run(u'Газрын төлбөр төлсөн байдлыг хүснэгтээр үзүүлэв:'+ u'                                                                            хэвлэсэн огноо: ' + str(date_now))
        font = run5.font
        font.name = 'Calibri'
        font.size = Pt(10)
        font.bold = False

        table = document.add_table(rows=1, cols=11)

        table.style = 'TableGrid'
        # table.style = 'LightShading-Accent1'
        # table.style = 'ColorfulList - Accent5'

        # table.style = 'TableNormal'
        table.allow_autofit = False
        table.autofit = False
        table.columns[0].width = Cm(0.1)
        table.columns[1].width = Cm(0.05)

        cell = table.rows[0].cells[0]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(u'№')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True

        cell = table.rows[0].cells[1]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(u'Он')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True

        cell = table.rows[0].cells[2]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(u'Төлбөр тооцсон хугацаа')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True

        cell = table.rows[0].cells[3]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(u'Төлөх төлбөр')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True

        cell = table.rows[0].cells[4]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(u'Төлсөн огноо')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True

        cell = table.rows[0].cells[5]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(u'Төлбөр төлсөн')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True

        cell = table.rows[0].cells[6]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(u'Төлбөрийн үлдэгдэл')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True

        cell = table.rows[0].cells[7]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(u'Алданги тооцсон')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True

        cell = table.rows[0].cells[8]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(u'Хоног тутам 0,1 хувиар')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True

        cell = table.rows[0].cells[9]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(u'Алдангийн хэмжээ')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True

        cell = table.rows[0].cells[10]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(u'Нийт төлөх төлбөрийн хэмжээ')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True

        surplus_last_year = 0

        begin_date = PluginUtils.convert_qt_date_to_python(self.begin_date.date())
        end_date = PluginUtils.convert_qt_date_to_python(self.end_date.date())

        sum_fee_contract = 0
        sum_paid = 0
        sum_balance = 0
        sum_total_fee = 0
        count = 0
        for payment_year in range(begin_date.year, end_date.year):

            amount_paid = self.session.query(func.sum(CtFeePayment.amount_paid)) \
                .filter(CtFeePayment.contract == fee.contract).filter(CtFeePayment.person == fee.person) \
                .filter(CtFeePayment.year_paid_for == payment_year).scalar()
            amount_count = self.session.query(CtFeePayment.date_paid) \
                .filter(CtFeePayment.contract == fee.contract).filter(CtFeePayment.person == fee.person) \
                .filter(CtFeePayment.year_paid_for == payment_year).count()
            date_paid = ''
            if amount_count != 0:
                amount = self.session.query(CtFeePayment) \
                    .filter(CtFeePayment.contract == fee.contract).filter(CtFeePayment.person == fee.person) \
                    .filter(CtFeePayment.year_paid_for == payment_year).all()
                for a in amount:
                    date_paid = str(a.date_paid)
            if amount_paid is None:
                amount_paid = 0

            fee_to_pay = self.__act_fee_to_pay_per_period(fee, date(payment_year, 1, 1), date(payment_year + 1, 1, 1))
            # print str(payment_year) + ' - ' + str(fee_to_pay) + ' - ' + str(amount_paid)
            if (amount_paid + surplus_last_year) - fee_to_pay > 0:
                surplus_last_year = (amount_paid + surplus_last_year) - fee_to_pay
            else:
                surplus_last_year = 0

            count = count + 1
            row_cells = table.add_row().cells
            # row_cells[0].text = str(count)
            cell = row_cells[0]
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(str(count))
            font = run.font
            font.name = 'Calibri'
            font.size = Pt(8)
            run.bold = True

            # row_cells[1].text = str(payment_year)
            cell = row_cells[1]
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(str(payment_year))
            font = run.font
            font.name = 'Calibri'
            font.size = Pt(8)
            run.bold = True

            # row_cells[2].text = self.payment_frequency_edit.text()
            cell = row_cells[2]
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(self.payment_frequency_edit.text())
            font = run.font
            font.name = 'Calibri'
            font.size = Pt(8)
            run.bold = True

            # row_cells[3].text = str(fee_to_pay)
            cell = row_cells[3]
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(str(fee_to_pay))
            font = run.font
            font.name = 'Calibri'
            font.size = Pt(8)
            run.bold = True
            sum_fee_contract = sum_fee_contract + fee_to_pay

            # row_cells[4].text = date_paid
            cell = row_cells[4]
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(date_paid)
            font = run.font
            font.name = 'Calibri'
            font.size = Pt(8)
            run.bold = True

            # row_cells[5].text = str(amount_paid)
            cell = row_cells[5]
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(str(amount_paid))
            font = run.font
            font.name = 'Calibri'
            font.size = Pt(8)
            run.bold = True
            sum_paid = sum_paid + amount_paid

            # row_cells[6].text = str(fee_to_pay - amount_paid)
            balance = fee_to_pay - amount_paid
            cell = row_cells[6]
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(str(balance))
            font = run.font
            font.name = 'Calibri'
            font.size = Pt(8)
            run.bold = True
            sum_balance = sum_balance + balance

            cell = row_cells[10]
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(str(balance))
            font = run.font
            font.name = 'Calibri'
            font.size = Pt(8)
            run.bold = True
            sum_total_fee = sum_total_fee + balance


        self.__set_column_width(table.columns[0], Cm(0.7))
        self.__set_column_width(table.columns[1], Cm(1))
        self.__set_column_width(table.columns[2], Cm(2.5))
        self.__set_column_width(table.columns[3], Cm(1.5))
        self.__set_column_width(table.columns[4], Cm(1.8))
        self.__set_column_width(table.columns[5], Cm(1.5))
        self.__set_column_width(table.columns[6], Cm(1.8))
        self.__set_column_width(table.columns[7], Cm(1.5))
        self.__set_column_width(table.columns[8], Cm(2.1))
        self.__set_column_width(table.columns[9], Cm(1.8))
        self.__set_column_width(table.columns[10], Cm(2.8))

        row_cells = table.add_row().cells
        cell = row_cells[3]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(str(sum_fee_contract))
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True

        cell = row_cells[5]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(str(sum_paid))
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True

        cell = row_cells[6]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(str(sum_balance))
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True

        cell = row_cells[10]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(str(sum_total_fee))
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True

        a = row_cells[0]
        b = row_cells[2]
        A = a.merge(b)

        cell = row_cells[0]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(u'Дүн')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(8)
        run.bold = True

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(u'1. Монгол Улсын Газрын тухай хуулийн 7 дугаар зүйлийн  7.1 дэх заалт "Газар эзэмшиж, ашиглаж байгаа иргэн, аж ахуйн нэгж, байгууллага зохих хууль тогтоомж, гэрээний дагуу газрын төлбөр төлнө" \n')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(10)
        font.bold = True

        run1 = p.add_run(u'2. Монгол Улсын Газрын тухай хуулийн 35 дугаар зүйлийн 35.3.1 дэх заалт "Газар эзэмших гэрээнд заасан нөхцөл, болзлыг биелүүлэх" \n')
        font = run1.font
        font.name = 'Calibri'
        font.size = Pt(10)
        font.bold = True

        run2 = p.add_run(u'3. Монгол Улсын Газрын тухай хуулийн 35 дугаар зүйлийн 35.3.3 дахь заалт "Газрын төлбөрийг тогтоосон хугацаанд төлөх " \n')
        font = run2.font
        font.name = 'Calibri'
        font.size = Pt(10)
        font.bold = True

        run3 = p.add_run(u'4. Монгол Улсын Газрын төлбөрийн тухай хуулийн 9 дүгээр зүйлийн 9.2 дахь заалт "Газрын төлбөр төлөгч нь газар эзэмшүүлэх, ашиглуулах гэрээнд өөрөөр заагаагүй бол жилийн төлбөрийг тэнцүү хэмжээгээр хуваан улирал бүрийн эхний сарын 24-ны дотор төлөх бөгөөд дараа улирлуудын төлбөрийг урьдчилан төлж болно." \n')
        font = run3.font
        font.name = 'Calibri'
        font.size = Pt(10)
        font.bold = True

        run4 = p.add_run(u'Дээрх заалтуудыг зөрчин газар эзэмшсэний төлбөр төлөөгүй тохиолдолд шүүхэд нэхэмжилж, Монгол Улсын Газрын тухай хуулийн 40 дүгээр зүйлийн 40.1.5 дахь заалт "Эрхийн гэрчилгээ эзэмшигч газрын төлбөрөө хугацаанд нь бүрэн төлөөгүй" тохиолдолд газар эзэмших эрхийн гэрчилгээг хүчингүй болгох арга хэмжээ авахыг үүгээр мэдэгдье. \n')
        font = run4.font
        font.name = 'Calibri'
        font.size = Pt(10)
        font.bold = False

        # document.add_paragraph('Intense quote', style='IntenseQuote')
        # document.add_page_break()
        fee_text = ''
        balance = sum_paid - sum_total_fee
        if sum_total_fee == 0:
            fee_text = u'ГАЗРЫН ТӨЛБӨРИЙН ҮЛДЭГДЭЛГҮЙ'
        elif sum_total_fee > 0:
            fee_text = u'ГАЗРЫН ТӨЛБӨРИЙН ҮЛДЭГДЭЛ '+ str(sum_total_fee) + u' ТӨГРӨГ'
        else:
            fee_text = u'ГАЗРЫН ТӨЛБӨРИЙН ИЛҮҮ ТӨЛӨЛТ ' + str(balance) + u' ТӨГРӨГ'
        text = str(begin_date.year) +' - '+ str(end_date.year) + u' оны хувьд газрын төлбөрийн тооцоо нийлэхэд ' + \
            fee_text + u' байна. '
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = p.add_run(text)
        font = p.font
        font.name = 'Calibri'
        font.size = Pt(12)
        font.bold = True

        report_settings = self.__admin_settings("set_report_parameter")

        land_office_name = report_settings[Constants.REPORT_LAND_OFFICE_NAME]
        bank_and_account = report_settings[Constants.REPORT_BANK] + u" банкны " + report_settings[
            Constants.REPORT_BANK_ACCOUNT]
        account_text = u'Газрын төлбөрийг ' + bank_and_account + u' тоот ' + land_office_name + u' -ын дансанд газрын төлбөр барагдуулах хуваарийн дагуу хугацаанд нь тушаана уу.'

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = p.add_run(account_text)
        font = p.font
        font.name = 'Calibri'
        font.size = Pt(8)
        font.bold = True

        user_name = QSettings().value(SettingsConstants.USER)
        print_officers = self.session.query(SetRole) \
            .filter(SetRole.user_name == user_name) \
            .filter(SetRole.is_active == True).one()
        position = print_officers.position_ref.name

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        off_name = print_officers.surname[:1]+'.'+print_officers.first_name
        officer_phone = ''
        if print_officers.phone:
            officer_phone = str(print_officers.phone)
        person_name = ''
        if len(info.person_register) == 10:
            person_name = info.name[:1] +'.'+ info.first_name
        else:
            if not info.contact_first_name:
                PluginUtils.show_message(self, self.tr('Contact name'), self.tr('Please contact person register!!!'))
                return
            person_name = info.contact_surname[:1] +'.'+ info.contact_first_name
        officer = p.add_run(u'                              АКТ ТОГТООСОН:                                           ' + off_name + '  / '+officer_phone+ '/' +'\n')
        font = officer.font
        font.name = 'Calibri'
        font.size = Pt(12)
        font.bold = True

        officer1 = p.add_run('                                                                                                                                                                               /'+position + '/ \n')
        font = officer1.font
        font.name = 'Calibri'
        font.size = Pt(6)
        font.bold = True

        officer2 = p.add_run(u'                              ТООЦОО НИЙЛСЭН:                                      ' + person_name + '\n')
        font = officer2.font
        font.name = 'Calibri'
        font.size = Pt(12)
        font.bold = True

        try:
            document.save('D:\demo.docx')
        except IOError, e:
            PluginUtils.show_error(self, self.tr("Out error"),
                                   self.tr("This file is already opened. Please close re-run"))

        try:
            QDesktopServices.openUrl(QUrl.fromLocalFile("D:/demo.docx"))
        except IOError, e:
            PluginUtils.show_error(self, self.tr("Out error"),
                                   self.tr("This file is already opened. Please close re-run"))

        return surplus_last_year

    def __admin_settings(self, table_name):

        session = SessionHandler().session_instance()
        lookup = {}
        l2_code = DatabaseUtils.working_l2_code()
        try:
            sql = "SELECT * FROM " + "s" + l2_code + ".{0};".format(table_name)
            result = session.execute(sql).fetchall()
            for row in result:
                lookup[row[0]] = row[1]

        except exc.SQLAlchemyError, e:
            PluginUtils.show_error(self, self.tr("SQL Error"), e.message)

        return lookup

    def __act_fee_to_pay_per_period(self, fee, period_begin, period_end):

        # Intersect contract duration with payment period
        sql = "select lower(daterange(contract_begin, contract_end, '[)') * daterange(:from, :to, '[)'))," \
              " upper(daterange(contract_begin, contract_end, '[)') * daterange(:from, :to, '[)')) " \
              "from ct_contract where contract_id = :contract_id"

        result = self.session.execute(sql, {'from': period_begin,
                                            'to': period_end,
                                            'contract_id': fee.contract})
        for row in result:
            effective_begin = row[0]
            effective_end = row[1]

        if effective_begin is None and effective_end is None:
            return 0

        # Intersect the effective payment period with the archived fees
        sql = "select upper(daterange(valid_from, valid_till, '[)') * daterange(:begin, :end, '[)')) - " \
              "lower(daterange(valid_from, valid_till, '[)') * daterange(:begin, :end, '[)')) as days, " \
              "fee_contract from ct_archived_fee where contract = :contract and person = :person"

        result = self.session.execute(sql, {'begin': effective_begin,
                                            'end': effective_end,
                                            'contract': fee.contract,
                                            'person': fee.person})
        fee_for_period = 0
        total_days = 0

        for row in result:
            days = row[0]
            if days is None:
                continue
            annual_fee = row[1]
            adjusted_fee = (annual_fee / 365) * days
            fee_for_period += adjusted_fee
            total_days += days

        effective_days = (effective_end - effective_begin).days

        if effective_days - total_days > 0:
            fee_for_period += (effective_days - total_days) * fee.fee_contract / 365

        return int(round(fee_for_period))